import bs4
import csv
import docx
import requests

print('\n ********COVID-19 informations**********')
print('1. District wise cases of karnataka state')
print('2. State wise cases')
print('3. Total cases,deaths,active cases in India')
option = int(input("Enter your option:"))
if option == 1:
    print('Googling........')         # display text while downloading the page.
    try:
                                      # getting request from the web page
        page = requests.get(
            "https://www.grainmart.in/news/covid-19-coronavirus-india-state-and-district-wise-tally/")
    
        soup = bs4.BeautifulSoup(page.content, 'html.parser')  # passing requested object to BeautifulSoup

    # From the table getting control of Karnataka state
        karnataka = soup.find_all(class_='skgm-districts')[1]
    # holding control of all rows of the table (all districts informaion under karnataka state)
        districts = karnataka.find_all(class_='skgm-tr')
        all_districts = []  # list to hold names of all districts
        active_cases = []  # list to hold active cases of all districts
        total_cases = []  # list to hold cases of all districts
        total_cured = []  # list to hold cured data of all districts
        total_deaths = []  # list to hold total deaths of all districts
        for item in districts:  # looping over each districts to collect all informations

            third_col = item.find_all(class_='skgm-td')[3]  # active cases
            cases1 = item.find_all(class_='skgm-td')[1]
            cases = cases1.find(class_='td-dc').getText()  # total cases
            cured1 = item.find_all(class_='skgm-td')[2]
            cured = cured1.find(class_='td-dr').getText()  # no. of cured people
            death1 = item.find_all(class_='skgm-td')[4]
            death = death1.find(class_='td-dd').getText()  # total deaths
            # appending the collected informations to the lists.
            all_districts.append(item.find_all(class_='skgm-td')[0].getText())
            active_cases.append(third_col.find(class_='td-da').getText())
            total_cases.append(cases)
            total_cured.append(cured)
            total_deaths.append(death)

        print('Successfully scraped data from web')

   # -------CONVERTING SCRAPED DATA TO CSV FORMAT-------------

    # Opening CSV file with write mode
        csvfile = open('district_table.csv', 'w', newline='')
        with csvfile:
            header = ['DISTRICTS', 'CASES', 'CURED', 'ACTIVE_CASES',
                  'DEATHS']  # Declaring header of the csv file
            writer = csv.DictWriter(csvfile, fieldnames=header)
            writer.writeheader()                 # writes the header to the csv file
            for i in range(31):  # loping over 30 districts to collect information from lists
                writer.writerow({                # writing a row to the csv file.
                    'DISTRICTS': all_districts[i],
                    'CASES': total_cases[i],
                    'CURED': total_cured[i],
                    'ACTIVE_CASES': active_cases[i],
                    'DEATHS': total_deaths[i]
                })
        print('Successfully converted to CSV format')

# -------CONVERTING CSV TO WORD DOCUMENT--------------
        doc = docx.Document()

        with open('district_table.csv', newline='') as f:  # opening csv file as f
            csv_reader = csv.reader(f)  # reading csv file
            csv_headers = next(csv_reader)               # header of csv file
            csv_cols = len(csv_headers)  # no.of columns

        # adding table to document with 2 blank rows below the heading(below the first row)
            table = doc.add_table(rows=2, cols=csv_cols)
        # variable hdr_cells has the control of cells of 0th row.
            hdr_cells = table.rows[0].cells

            for i in range(csv_cols):
                hdr_cells[i].text = csv_headers[i]       # Adding headings

            for row in csv_reader:
            # adding one row each time to hold the values.
                row_cells = table.add_row().cells
                for i in range(csv_cols):
                # adding information to each cell in that row.
                    row_cells[i].text = row[i]

        doc.add_page_break()
        doc.save("district_table.docx")                # saving word document.

        print("Successfully converted CSV to Word document.")


    except:                                            # error message
        print('An error occured while scraping!! Check your internet connection.\n Close CSV file and Word Document,if it is open!')

elif option == 2:
    try:
        print('Googling........')  # display text while downloading the page.
    
        page = requests.get(
                "https://www.grainmart.in/news/covid-19-coronavirus-india-state-and-district-wise-tally/")
        soup = bs4.BeautifulSoup(page.content, 'html.parser')
        #STATES=soup.find_all(class_='skgm-th')
        all_states1 = []    #list to hold all state names
        active_cases1 = []  # list to hold active cases
        total_cases1 = []  # list to hold cases of all states
        total_cured1 = []  # list to hold cured data of all states
        total_deaths1 = []
        for i in range(35):
            state=soup.find_all(class_='skgm-states')[i]     # row consisting of one state information
            all_states1.append(state.find_all(class_='skgm-td')[0].getText()) # appending state name
            case1=state.find_all(class_='skgm-td')[1]
            total_cases1.append(case1.find(class_='td-sc').getText()) #appending total cases
            cured1=state.find_all(class_='skgm-td')[2]
            total_cured1.append(cured1.find(class_='td-sr').getText()) #appending total cured data
            active1=state.find_all(class_='skgm-td')[3]
            active_cases1.append(active1.find(class_='td-sa').getText()) #appending active cases data
            death1=state.find_all(class_='skgm-td')[4]
            total_deaths1.append(death1.find(class_='td-sd').getText()) #appending info. of total death

        print('Successfully scraped data from web')

    # -------CONVERTING SCRAPED DATA TO CSV FORMAT-------------

    # Opening CSV file with write mode
        csvfile = open('state_table.csv', 'w', newline='')
        with csvfile:
            header = ['STATES', 'CASES', 'CURED', 'ACTIVE_CASES',
                  'DEATHS']  # Declaring header of the csv file
            writer = csv.DictWriter(csvfile, fieldnames=header)
            writer.writeheader()                 # writes the header to the csv file
            for i in range(35):  # loping over 35 states to collect information from lists
                writer.writerow({                # writing a row to the csv file.
                    'STATES': all_states1[i],
                    'CASES': total_cases1[i],
                    'CURED': total_cured1[i],
                    'ACTIVE_CASES': active_cases1[i],
                    'DEATHS': total_deaths1[i]
                })
        print('Successfully converted to CSV format')

        # -------CONVERTING CSV TO WORD DOCUMENT--------------
        doc = docx.Document()

        with open('state_table.csv', newline='') as f:  # opening csv file as f
            csv_reader = csv.reader(f)  # reading csv file
            csv_headers = next(csv_reader)               # header of csv file
            csv_cols = len(csv_headers)  # no.of columns

        # adding table to document with 2 blank rows below the heading(below the first row)
            table = doc.add_table(rows=2, cols=csv_cols)
        # variable hdr_cells has the control of cells of 0th row.
            hdr_cells = table.rows[0].cells

            for i in range(csv_cols):
                hdr_cells[i].text = csv_headers[i]       # Adding headings

            for row in csv_reader:
            # adding one row each time to hold the values.
                row_cells = table.add_row().cells
                for i in range(csv_cols):
                # adding information to each cell in that row.
                    row_cells[i].text = row[i]

        doc.add_page_break()
        doc.save("state_table.docx")                # saving word document.

        print("Successfully converted CSV to Word document.")

    except:                    # error message,if no wifi connection to load web link
        print('An error occured while scraping!! Check your internet connection.\n Close CSV file and Word Document,if it is open!')

elif option == 3:
    print('Googling....')
    try:
        page = requests.get(
                "https://www.grainmart.in/news/covid-19-coronavirus-india-state-and-district-wise-tally/")
        soup = bs4.BeautifulSoup(page.content, 'html.parser')
        TOTAL=soup.find_all(class_='skgm-th')[0]
        heading=TOTAL.find_all(class_='skgm-td')[0].getText()  # heading of the row
        CASES1=TOTAL.find_all(class_='skgm-td')[1]
        CASE=CASES1.find(class_='td-tc').getText()  #total cases in India
        CURED1=TOTAL.find_all(class_='skgm-td')[2]
        CURED=CURED1.find(class_='td-tr').getText() # no. of cured people
        ACTIVE1=TOTAL.find_all(class_='skgm-td')[3]
        ACTIVE=ACTIVE1.find(class_='td-ta').getText()  # no. of active cases in India
        DEATH1=TOTAL.find_all(class_='skgm-td')[4]
        DEATH=DEATH1.find(class_='td-td').getText()  #total deaths

        print('Successfully scraped data from web')

        # -------CONVERTING SCRAPED DATA TO CSV FORMAT-------------

    # Opening CSV file with write mode
        csvfile = open('TOTAL.csv', 'w', newline='')
        with csvfile:
            header = ['TOTAL', 'CASES', 'CURED', 'ACTIVE_CASES',
                  'DEATHS']  # Declaring header of the csv file
            writer = csv.DictWriter(csvfile, fieldnames=header)
            writer.writeheader()                 # writes the header to the csv file
             
            writer.writerow({                # writing a row to the csv file.
                'TOTAL': heading,
                'CASES': CASE,
                'CURED': CURED,
                'ACTIVE_CASES': ACTIVE,
                'DEATHS': DEATH
            })
        print('Successfully converted to CSV format')

        # -------CONVERTING CSV TO WORD DOCUMENT--------------
        doc = docx.Document()

        with open('TOTAL.csv', newline='') as f:  # opening csv file as f
            csv_reader = csv.reader(f)            # reading csv file
            csv_headers = next(csv_reader)        # header of csv file
            csv_cols = len(csv_headers)           # no.of columns

        # adding table to document with 2 blank rows below the heading(below the first row)
            table = doc.add_table(rows=2, cols=csv_cols)
        # variable hdr_cells has the control of cells of 0th row.
            hdr_cells = table.rows[0].cells

            for i in range(csv_cols):
                hdr_cells[i].text = csv_headers[i]       # Adding headings

            for row in csv_reader:
            # adding one row each time to hold the values.
                row_cells = table.add_row().cells
                for i in range(csv_cols):
                # adding information to each cell in that row.
                    row_cells[i].text = row[i]

        doc.add_page_break()
        doc.save("TOTAL.docx")                # saving word document.

        print("Successfully converted CSV to Word document.")


    except:
        print('An error occured while scraping!! Check your internet connection.\n Close CSV file and Word Document,if it is open!')


else:                     # If option is not equal to 1 or 2 or 3.
    print('INVALID OPTION !!!')











    



